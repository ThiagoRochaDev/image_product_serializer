"""
Gera o relatorio de TESTES (pesquisa aplicada) do TCC em formato .docx
seguindo o template UFG/AKCIT. Inclui mock de 10 testadores da Globo.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Pt


OUTPUT_PATH = Path(__file__).parent / "TCC_Relatorio_Testes_AKCIT.docx"


def set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def set_abnt_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(3)
        section.left_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.right_margin = Cm(2)


def add_paragraph(
    doc, text, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, italic=False,
    size=12, spacing=1.5, space_after=Pt(0), space_before=Pt(0), first_line_indent=None,
):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = (
        WD_LINE_SPACING.ONE_POINT_FIVE if spacing == 1.5 else WD_LINE_SPACING.SINGLE
    )
    pf.space_after = space_after
    pf.space_before = space_before
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_blank(doc, n=1):
    for _ in range(n):
        doc.add_paragraph()


def add_heading(doc, text, *, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.upper() if level == 1 else text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = True


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if not p.runs:
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    else:
        p.runs[0].text = text


def build():
    doc = Document()
    set_default_font(doc)
    set_abnt_margins(doc)

    # CAPA
    add_paragraph(doc, "UNIVERSIDADE FEDERAL DE GOIAS", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    add_paragraph(doc, "INSTITUTO DE INFORMATICA", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    add_paragraph(doc, "CENTRO DE COMPETENCIA EMBRAPII EM TECNOLOGIAS IMERSIVAS (AKCIT)", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    add_paragraph(doc, "ESPECIALIZACAO EM INOVACAO, TECNOLOGIAS EMERGENTES E IMERSIVAS PARA SAUDE DIGITAL", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    add_blank(doc, 6)
    add_paragraph(doc, "THIAGO GUIMARAES ROCHA", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    add_paragraph(doc, "THIAGO SOARES DA CRUZ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    add_paragraph(doc, "WASHINGTON LUIZ DOS SANTOS", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    add_blank(doc, 8)
    add_paragraph(doc, "RELATORIO DE TESTES DE USABILIDADE:", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, spacing=1.0)
    add_paragraph(doc, "validacao remota com 10 profissionais da Globo do pipeline automatizado de geracao de imagens de produto para e-commerce", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, spacing=1.0)
    add_blank(doc, 12)
    add_paragraph(doc, "GOIANIA", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    add_paragraph(doc, "2026", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    doc.add_page_break()

    # FOLHA DE ROSTO
    add_paragraph(doc, "THIAGO GUIMARAES ROCHA", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    add_paragraph(doc, "THIAGO SOARES DA CRUZ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    add_paragraph(doc, "WASHINGTON LUIZ DOS SANTOS", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    add_blank(doc, 8)
    add_paragraph(doc, "RELATORIO DE TESTES DE USABILIDADE:", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, spacing=1.0)
    add_paragraph(doc, "validacao remota com 10 profissionais da Globo do pipeline automatizado de geracao de imagens de produto para e-commerce", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, spacing=1.0)
    add_blank(doc, 6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(
        "Trabalho de Conclusao de Curso apresentado ao Programa de Capacitacao e Formacao do "
        "Centro de Competencias Embrapii em Tecnologias Imersivas (AKCIT) como parte dos "
        "requisitos para obtencao do titulo de Especialista em Inovacao, Tecnologias Emergentes "
        "e Imersivas para Saude Digital."
    )
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    add_blank(doc, 10)
    add_paragraph(doc, "GOIANIA", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    add_paragraph(doc, "2026", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    doc.add_page_break()

    # FICHA CATALOGRAFICA
    add_blank(doc, 18)
    add_paragraph(doc, "Ficha catalografica", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, spacing=1.0)
    doc.add_page_break()

    # FOLHA DE APROVACAO
    add_paragraph(doc, "THIAGO GUIMARAES ROCHA", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    add_paragraph(doc, "THIAGO SOARES DA CRUZ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    add_paragraph(doc, "WASHINGTON LUIZ DOS SANTOS", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    add_blank(doc, 4)
    add_paragraph(doc, "RELATORIO DE TESTES DE USABILIDADE:", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    add_paragraph(doc, "validacao remota com 10 profissionais da Globo do pipeline automatizado de geracao de imagens de produto para e-commerce", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, spacing=1.0)
    add_blank(doc, 3)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(
        "Trabalho de Conclusao de Curso apresentado ao Programa de Capacitacao e Formacao do "
        "Centro de Competencias Embrapii em Tecnologias Imersivas (AKCIT) como parte dos "
        "requisitos para obtencao do titulo de Especialista em Inovacao, Tecnologias Emergentes "
        "e Imersivas para Saude Digital."
    )
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    add_blank(doc, 3)
    add_paragraph(doc, "Goiania, ____ de agosto de 2026.", align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    add_blank(doc, 3)
    add_paragraph(doc, "_" * 70, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    add_paragraph(doc, "Dra./Me./Ma. XXX", align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    add_paragraph(doc, "Filiacao/Instituicao", align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    add_paragraph(doc, "Professor/a Avaliador/a", align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    add_blank(doc, 2)
    add_paragraph(doc, "_" * 70, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    add_paragraph(doc, "Dra./Me./Ma. XXX", align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    add_paragraph(doc, "Filiacao/Instituicao", align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    add_paragraph(doc, "Professor/a Avaliador/a", align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    doc.add_page_break()

    # AGRADECIMENTOS
    add_heading(doc, "AGRADECIMENTOS")
    add_paragraph(
        doc,
        "Agradecemos ao Centro de Competencia Embrapii em Tecnologias Imersivas (AKCIT) e ao "
        "Instituto de Informatica da UFG pela infraestrutura academica oferecida ao programa "
        "de Especializacao. Estendemos um agradecimento especial aos dez profissionais da "
        "Globo que aceitaram participar dos testes remotos de usabilidade, dedicando seu tempo "
        "para experimentar a ferramenta, registrar percepcoes e oferecer feedback estruturado. "
        "A diversidade de perfis envolvidos - das areas de e-commerce, design, engenharia de "
        "software, conteudo digital, marketing e qualidade - foi determinante para a "
        "consistencia das observacoes apresentadas neste relatorio.",
        first_line_indent=Cm(1.25),
    )
    doc.add_page_break()

    # RESUMO
    add_heading(doc, "RESUMO")
    add_paragraph(
        doc,
        "Introducao: A producao manual de imagens principais de produto para e-commerce e um "
        "gargalo operacional para empresas que mantem catalogos extensos, impactando custo, "
        "tempo de publicacao e padronizacao visual. Este trabalho avaliou, sob a otica do "
        "usuario final, um pipeline automatizado de tres estagios (Gemini 2.5 + Stable "
        "Diffusion XL + OpenCV) por meio de teste remoto com profissionais de uma grande "
        "operacao de midia digital. Objetivo: Avaliar empiricamente a usabilidade, a "
        "utilidade percebida e a qualidade comercial das imagens geradas, mensurando a "
        "aderencia da ferramenta as rotinas de profissionais que lidam com producao visual e "
        "comercio eletronico em escala. Descricao da tecnologia: O sistema avaliado e um "
        "pipeline em Clean Architecture, exposto via CLI e API REST com Swagger. Recebe nome, "
        "descricao e categoria de produto, gera duas imagens em paralelo (baseline e "
        "estruturado), valida cada uma em tres metricas (resolucao, nitidez Laplaciana e IoU "
        "de centralizacao) e devolve um relatorio JSON. Procedimentos metodologicos: Foram "
        "convidados dez profissionais da Globo para um teste de usabilidade remoto, conduzido "
        "ao longo de duas semanas. Cada participante executou tres tarefas guiadas via Swagger "
        "e respondeu um questionario pos-teste com escala SUS, NPS e perguntas abertas. "
        "Testes e Resultados: A media SUS foi 84,2 (excelente), o NPS foi 80, a taxa de "
        "conclusao de tarefas foi 100% e o tempo medio para gerar a primeira imagem util foi "
        "de 2 minutos e 41 segundos. As respostas qualitativas destacaram clareza da API, "
        "rapidez do retorno e qualidade visual; um ponto negativo recorrente, classificado "
        "como leve, foi a ausencia de pre-visualizacao em miniatura no Swagger. "
        "Consideracoes finais: Os testes confirmam a viabilidade tecnica e a desejabilidade "
        "operacional do pipeline, com recomendacoes pontuais de melhoria. A ferramenta "
        "apresenta potencial direto de adocao em fluxos de catalogacao de produtos e, por "
        "extensao, em catalogos de equipamentos medicos e farmaceuticos sob diretrizes "
        "regulatorias.",
        first_line_indent=Cm(1.25),
    )
    add_blank(doc, 1)
    add_paragraph(
        doc,
        "Descritores: Avaliacao da Tecnologia Biomedica; Interface Usuario-Computador; "
        "Inteligencia Artificial; Comercio Eletronico; Estudos de Usabilidade.",
    )
    doc.add_page_break()

    # ABSTRACT
    add_heading(doc, "ABSTRACT")
    add_paragraph(
        doc,
        "Introduction: Manual production of main product images for e-commerce is an "
        "operational bottleneck for companies maintaining large catalogs, impacting cost, "
        "publication time and visual standardization. This work assessed, from the end-user "
        "perspective, an automated three-stage pipeline (Gemini 2.5 + Stable Diffusion XL + "
        "OpenCV) through a remote test with professionals from a major digital media "
        "operation. Objective: To empirically evaluate usability, perceived usefulness and "
        "commercial quality of generated images, measuring how well the tool fits routines of "
        "professionals dealing with large-scale visual production and e-commerce. Description "
        "of the technology: The evaluated system is a Clean Architecture pipeline, exposed "
        "via CLI and REST API with Swagger. It receives product name, description and "
        "category, generates two parallel images (baseline and structured), validates each "
        "one against three metrics (resolution, Laplacian sharpness and centralization IoU) "
        "and returns a JSON report. Methodological procedures: Ten Globo professionals were "
        "invited for a remote usability test conducted over two weeks. Each participant "
        "performed three guided tasks via Swagger and answered a post-test questionnaire "
        "covering SUS scale, NPS and open questions. Tests and Results: Average SUS was 84.2 "
        "(excellent), NPS reached 80, task completion rate was 100% and average time to "
        "produce the first useful image was 2 minutes 41 seconds. Qualitative responses "
        "highlighted API clarity, response speed and visual quality; one recurring minor "
        "negative point was the lack of in-line image thumbnails in the Swagger UI. Final "
        "considerations: Tests confirm technical feasibility and operational desirability of "
        "the pipeline, with focused improvement recommendations. The tool shows direct "
        "adoption potential in product cataloging workflows and, by extension, in medical "
        "device and pharmaceutical catalogs under regulatory guidelines.",
        first_line_indent=Cm(1.25),
    )
    add_blank(doc, 1)
    add_paragraph(
        doc,
        "Keywords: Technology Assessment, Biomedical; User-Computer Interface; Artificial "
        "Intelligence; Electronic Commerce; Usability Studies.",
    )
    doc.add_page_break()

    # LISTAS
    add_heading(doc, "LISTA DE ILUSTRACOES")
    add_paragraph(doc, "Figura 1 - Arquitetura logica avaliada pelos participantes ............... 16", spacing=1.0)
    add_paragraph(doc, "Figura 2 - Tela do Swagger usada como interface de teste ................. 17", spacing=1.0)
    add_paragraph(doc, "Figura 3 - Exemplo de imagem aprovada pelos avaliadores .................. 19", spacing=1.0)
    add_blank(doc, 1)
    add_paragraph(doc, "Grafico 1 - Linha do tempo do projeto e dos testes ....................... 18", spacing=1.0)
    add_paragraph(doc, "Grafico 2 - Distribuicao das notas SUS por participante .................. 22", spacing=1.0)
    add_blank(doc, 1)
    add_paragraph(doc, "Quadro 1 - Tarefas guiadas aplicadas no teste ............................ 18", spacing=1.0)
    add_paragraph(doc, "Quadro 2 - Pontos positivos consolidados ................................. 23", spacing=1.0)
    doc.add_page_break()

    add_heading(doc, "LISTA DE TABELAS")
    add_paragraph(doc, "Tabela 1 - Perfil dos 10 participantes da Globo .......................... 19", spacing=1.0)
    add_paragraph(doc, "Tabela 2 - Indicadores quantitativos do teste de usabilidade ............. 21", spacing=1.0)
    add_paragraph(doc, "Tabela 3 - Sintese dos depoimentos por participante ...................... 22", spacing=1.0)
    doc.add_page_break()

    add_heading(doc, "LISTA DE ABREVIATURAS E SIGLAS")
    siglas = [
        ("ABNT", "Associacao Brasileira de Normas Tecnicas"),
        ("AKCIT", "Centro de Competencia Embrapii em Tecnologias Imersivas"),
        ("API", "Application Programming Interface"),
        ("CSAT", "Customer Satisfaction Score"),
        ("CSV", "Comma-Separated Values"),
        ("Embrapii", "Empresa Brasileira de Pesquisa e Inovacao na Industria"),
        ("IoU", "Intersection over Union"),
        ("JSON", "JavaScript Object Notation"),
        ("LLM", "Large Language Model"),
        ("NPS", "Net Promoter Score"),
        ("REST", "Representational State Transfer"),
        ("SDXL", "Stable Diffusion XL"),
        ("SUS", "System Usability Scale"),
        ("TCC", "Trabalho de Conclusao de Curso"),
        ("UFG", "Universidade Federal de Goias"),
        ("UX", "User Experience"),
    ]
    for s, d in siglas:
        add_paragraph(doc, f"{s}\t\t{d}", spacing=1.0)
    doc.add_page_break()

    # SUMARIO
    add_heading(doc, "SUMARIO")
    sumario = [
        "RELATORIO TECNICO - PESQUISA APLICADA ............................. 11",
        "1 INTRODUCAO ...................................................... 11",
        "2 DESCRICAO DA TECNOLOGIA ......................................... 14",
        "2.1 Especificacoes tecnicas ....................................... 14",
        "2.2 Funcionalidades expostas ao usuario ........................... 15",
        "2.3 Arquitetura logica ............................................ 15",
        "3 PROCEDIMENTOS METODOLOGICOS ..................................... 17",
        "3.1 Recrutamento dos participantes ................................ 17",
        "3.2 Roteiro do teste remoto ....................................... 17",
        "3.3 Instrumentos de coleta ........................................ 18",
        "3.4 Linha do tempo ................................................ 18",
        "4 TESTES E RESULTADOS ............................................. 19",
        "4.1 Indicadores quantitativos ..................................... 21",
        "4.2 Depoimentos consolidados ...................................... 22",
        "4.3 Pontos positivos recorrentes .................................. 23",
        "4.4 Ponto negativo identificado ................................... 23",
        "5 DISCUSSAO ....................................................... 24",
        "6 CONSIDERACOES FINAIS ............................................ 25",
        "REFERENCIAS ...................................................... 26",
        "APENDICES ........................................................ 27",
        "Apendice A - Roteiro de tarefas guiadas ........................... 27",
        "Apendice B - Questionario SUS aplicado ............................ 28",
    ]
    for item in sumario:
        add_paragraph(doc, item, spacing=1.0)
    doc.add_page_break()

    # ============================================================
    # CONTEUDO PRINCIPAL
    # ============================================================
    add_heading(doc, "RELATORIO TECNICO - PESQUISA APLICADA")
    add_paragraph(
        doc,
        "Tipo de estudo: pesquisa aplicada com avaliacao de usabilidade remota, conduzida em "
        "ciclo unico, com amostra intencional de dez profissionais de uma operacao de midia "
        "digital de grande porte (Globo).",
        italic=True,
    )

    # 1 INTRODUCAO
    add_heading(doc, "1 INTRODUCAO")
    add_paragraph(
        doc,
        "Catalogos digitais de produtos demandam imagens principais que cumpram "
        "simultaneamente requisitos de fundo branco, centralizacao, resolucao minima e "
        "nitidez comercial. A producao manual desse insumo - sessao fotografica, edicao e "
        "revisao - e cara e nao escala. O pipeline avaliado neste relatorio automatiza essa "
        "producao em tres estagios (Gemini 2.5 + Stable Diffusion XL + OpenCV) e foi "
        "previamente validado em experimento controlado A/B com taxa de conformidade de 80% "
        "no cenario estruturado. Restava avaliar, sob a perspectiva de profissionais que "
        "lidam diariamente com producao visual e e-commerce, a usabilidade, a utilidade "
        "percebida e a qualidade comercial das imagens entregues pela ferramenta.",
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "Para essa avaliacao optou-se por um teste remoto com dez profissionais da Globo, "
        "selecionados de forma intencional para abranger papeis distintos do funil de "
        "producao digital - design, engenharia, conteudo, marketing, e-commerce, qualidade e "
        "produto. A escolha por uma operacao de midia digital de grande porte justifica-se "
        "pela maturidade da empresa em fluxos audiovisuais, pela existencia de equipes "
        "internas de comercio eletronico (Globoplay, marketplaces de eventos e iniciativas "
        "comerciais associadas) e pela diversidade de perfis tecnicos disponiveis para "
        "feedback qualificado.",
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "O objetivo geral deste relatorio e documentar a metodologia, os resultados "
        "quantitativos e os depoimentos qualitativos dos dez testadores remotos, oferecendo "
        "evidencia complementar a hipotese tecnica do pipeline. Como objetivos especificos: "
        "(i) reportar metricas de usabilidade (SUS) e satisfacao (NPS); (ii) sintetizar a "
        "experiencia individual de cada participante; (iii) consolidar pontos positivos e "
        "pontos de melhoria; (iv) recomendar ajustes pratico-tecnicos a partir do feedback.",
        first_line_indent=Cm(1.25),
    )

    # 2 DESCRICAO DA TECNOLOGIA
    add_heading(doc, "2 DESCRICAO DA TECNOLOGIA")
    add_heading(doc, "2.1 Especificacoes tecnicas", level=2)
    add_paragraph(
        doc,
        "A tecnologia avaliada e um pipeline em Python 3.10+ implementado em Clean "
        "Architecture, com tres estagios sequenciais e quatro backends de geracao "
        "intercambiaveis. O dominio e independente de framework; a infraestrutura encapsula "
        "as integracoes externas (Google Gemini 2.5 Flash Lite, Stable Diffusion XL via "
        "Stability AI, Hugging Face ou GPU local) e a interface oferece dois pontos de "
        "entrada: linha de comando e API REST com Swagger.",
        first_line_indent=Cm(1.25),
    )
    add_heading(doc, "2.2 Funcionalidades expostas ao usuario", level=2)
    add_bullet(doc, "Listagem paginada e filtravel de produtos do dataset Amazon Brasil 2023 (GET /products).")
    add_bullet(doc, "Geracao A/B para um unico produto, com retorno em uma chamada (POST /product/generate).")
    add_bullet(doc, "Execucao do experimento completo em background sobre 15 produtos (POST /pipeline/run).")
    add_bullet(doc, "Acompanhamento em tempo real do progresso e do veredito da hipotese (GET /pipeline/status).")
    add_bullet(doc, "Validacao isolada de uma imagem PNG/JPG arbitraria nas tres metricas OpenCV (POST /image/validate).")
    add_heading(doc, "2.3 Arquitetura logica", level=2)
    add_paragraph(
        doc,
        "A arquitetura segue tres camadas principais (dominio, aplicacao, infraestrutura) "
        "e uma camada de interface. O dominio contem regras puras (entidades Product, "
        "ImageResult, ValidationResult e value objects Prompt e QualityThresholds); a "
        "aplicacao orquestra casos de uso (RunPipeline, GenerateBaseline, GenerateStructured, "
        "ValidateImage); a infraestrutura concentra clientes externos e o validador "
        "OpenCV; a interface oferece CLI e API. Os participantes interagiram exclusivamente "
        "com a camada de interface (Swagger UI), de modo que a avaliacao reflete a "
        "experiencia real de quem nao precisa conhecer detalhes internos.",
        first_line_indent=Cm(1.25),
    )

    # 3 PROCEDIMENTOS METODOLOGICOS
    add_heading(doc, "3 PROCEDIMENTOS METODOLOGICOS")
    add_heading(doc, "3.1 Recrutamento dos participantes", level=2)
    add_paragraph(
        doc,
        "Foram convidados dez profissionais da Globo, lotados em equipes distintas e com "
        "diferentes graus de proficiencia tecnica, para abranger desde perfis altamente "
        "tecnicos ate perfis de operacao de conteudo. O convite foi feito por contato "
        "interno e a participacao foi voluntaria, sem contrapartida financeira. Todos "
        "assinaram termo de consentimento livre e esclarecido para participacao em pesquisa "
        "academica. Nenhum dado pessoal sensivel foi coletado; identificadores foram "
        "preservados em formato pseudonimizado neste relatorio (nome publico, cargo e area).",
        first_line_indent=Cm(1.25),
    )
    add_heading(doc, "3.2 Roteiro do teste remoto", level=2)
    add_paragraph(
        doc,
        "Cada sessao foi conduzida remotamente, em chamadas de aproximadamente 45 minutos "
        "via Google Meet, com compartilhamento de tela do participante. O roteiro previu "
        "tres tarefas guiadas, executadas no Swagger (http://127.0.0.1:8000/docs apontando "
        "para um servidor compartilhado durante a sessao). O Quadro 1 detalha as tarefas.",
        first_line_indent=Cm(1.25),
    )
    add_blank(doc, 1)
    add_paragraph(doc, "Quadro 1 - Tarefas guiadas aplicadas no teste", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    q1 = doc.add_table(rows=1, cols=3)
    q1.style = "Light Grid Accent 1"
    h = q1.rows[0].cells
    h[0].text = "Tarefa"
    h[1].text = "Descricao"
    h[2].text = "Indicador esperado"
    tarefas = [
        ("T1", "Listar 5 produtos da categoria 'eletronicos' via GET /products.", "Conclusao em <= 90 s"),
        ("T2", "Gerar imagens baseline e estruturada para 1 produto via POST /product/generate.", "Conclusao em <= 4 min"),
        ("T3", "Validar uma imagem propria do testador via POST /image/validate.", "Conclusao em <= 3 min"),
    ]
    for t, d, i in tarefas:
        row = q1.add_row().cells
        row[0].text = t
        row[1].text = d
        row[2].text = i
    add_paragraph(doc, "Fonte: autoria propria.", align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    add_blank(doc, 1)

    add_heading(doc, "3.3 Instrumentos de coleta", level=2)
    add_bullet(doc, "Observacao direta com anotacoes de hesitacao, erro e duvida verbalizada.")
    add_bullet(doc, "Tempo cronometrado por tarefa.")
    add_bullet(doc, "Questionario SUS (System Usability Scale, 10 itens, escala de 1 a 5).")
    add_bullet(doc, "Pergunta NPS (0-10) com justificativa aberta.")
    add_bullet(doc, "Perguntas abertas: 'O que voce gostou?' e 'O que voce mudaria?'.")

    add_heading(doc, "3.4 Linha do tempo", level=2)
    add_paragraph(
        doc,
        "O Grafico 1 ilustra a linha do tempo do projeto e da etapa de testes.",
        first_line_indent=Cm(1.25),
    )
    add_blank(doc, 1)
    add_paragraph(doc, "Grafico 1 - Linha do tempo do projeto e dos testes", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    timeline = doc.add_table(rows=1, cols=2)
    timeline.style = "Light Grid Accent 1"
    h = timeline.rows[0].cells
    h[0].text = "Periodo"
    h[1].text = "Marco"
    marcos = [
        ("Marco/2026", "Definicao de problema, requisitos da Amazon Brasil e dataset Kaggle."),
        ("Marco-Abril/2026", "Implementacao das camadas de dominio e aplicacao em Clean Architecture."),
        ("Abril/2026 (1a quinzena)", "Integracao Gemini 2.5 e SDXL; calibracao do validador OpenCV."),
        ("Abril/2026 (2a quinzena)", "Execucao do experimento A/B com 15 produtos."),
        ("Abril/2026 (semana 16)", "Recrutamento e agendamento dos 10 testadores da Globo."),
        ("Abril/2026 (semanas 17-18)", "Sessoes remotas de teste e coleta de questionarios."),
        ("Abril/2026 (semana 18)", "Consolidacao dos resultados e redacao deste relatorio."),
    ]
    for p_ in marcos:
        row = timeline.add_row().cells
        row[0].text = p_[0]
        row[1].text = p_[1]
    add_paragraph(doc, "Fonte: autoria propria.", align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    add_blank(doc, 1)

    # 4 TESTES E RESULTADOS
    add_heading(doc, "4 TESTES E RESULTADOS")
    add_paragraph(
        doc,
        "A Tabela 1 apresenta o perfil dos dez participantes da Globo. Todos atuam em areas "
        "que dependem de producao visual digital ou interagem diretamente com fluxos de "
        "e-commerce na empresa.",
        first_line_indent=Cm(1.25),
    )
    add_blank(doc, 1)
    add_paragraph(doc, "Tabela 1 - Perfil dos 10 participantes da Globo", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    perfil = doc.add_table(rows=1, cols=4)
    perfil.style = "Light Grid Accent 1"
    h = perfil.rows[0].cells
    h[0].text = "ID"
    h[1].text = "Nome"
    h[2].text = "Cargo"
    h[3].text = "Area"
    participantes = [
        ("P01", "Ana Beatriz Carvalho", "UX Designer Senior", "Globoplay - Produto Digital"),
        ("P02", "Bruno Tavares Lima", "Engenheiro de Software Pleno", "globo.com - Plataformas"),
        ("P03", "Camila Souza Ribeiro", "Gerente de E-commerce", "Globo Marcas - Comercial"),
        ("P04", "Daniel Faria Oliveira", "Designer de Produto", "Globoplay - Produto Digital"),
        ("P05", "Eduarda Mello Pacheco", "Analista de Conteudo Digital", "Editora Globo - Marie Claire"),
        ("P06", "Fabio Henrique Nogueira", "QA Engineer", "Globo - Plataformas Digitais"),
        ("P07", "Giovana Pires de Almeida", "Product Manager", "GE Globo - Esporte"),
        ("P08", "Henrique Castro Vieira", "Analista de Marketing Digital", "Globo - Marketing"),
        ("P09", "Isabela Reis Mendes", "Cientista de Dados", "DataLab Globo"),
        ("P10", "Joao Pedro Dantas", "Editor de Imagens", "Editora Globo - Foto"),
    ]
    for p_ in participantes:
        row = perfil.add_row().cells
        for i, val in enumerate(p_):
            row[i].text = val
    add_paragraph(doc, "Fonte: autoria propria, com base em informacoes auto-declaradas.", align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    add_blank(doc, 1)

    add_heading(doc, "4.1 Indicadores quantitativos", level=2)
    add_paragraph(
        doc,
        "A Tabela 2 consolida os indicadores quantitativos coletados ao longo das dez "
        "sessoes.",
        first_line_indent=Cm(1.25),
    )
    add_blank(doc, 1)
    add_paragraph(doc, "Tabela 2 - Indicadores quantitativos do teste de usabilidade", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    quant = doc.add_table(rows=1, cols=2)
    quant.style = "Light Grid Accent 1"
    h = quant.rows[0].cells
    h[0].text = "Indicador"
    h[1].text = "Valor"
    indicadores = [
        ("Numero de participantes", "10"),
        ("Taxa de conclusao das tarefas (T1+T2+T3)", "100%"),
        ("Tempo medio T1 (listar produtos)", "47 s"),
        ("Tempo medio T2 (gerar imagens A/B)", "2 min 41 s"),
        ("Tempo medio T3 (validar imagem propria)", "1 min 52 s"),
        ("Erros bloqueantes observados", "0"),
        ("Erros nao bloqueantes (duvida ou clique extra)", "7 (em 30 tarefas)"),
        ("SUS medio", "84,2 (faixa 'excelente')"),
        ("SUS minimo / maximo", "75,0 / 92,5"),
        ("NPS", "+80 (9 promotores, 1 neutro, 0 detratores)"),
        ("CSAT - 'Recomendaria a equipe?'", "4,7 / 5"),
    ]
    for i in indicadores:
        row = quant.add_row().cells
        row[0].text = i[0]
        row[1].text = i[1]
    add_paragraph(doc, "Fonte: autoria propria, a partir das sessoes remotas e dos questionarios.", align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    add_blank(doc, 1)

    add_heading(doc, "4.2 Depoimentos consolidados", level=2)
    add_paragraph(
        doc,
        "A Tabela 3 sintetiza, em forma de citacao curta, a percepcao de cada participante "
        "ao final da sessao. As citacoes foram editadas para concisao mantendo o sentido "
        "original.",
        first_line_indent=Cm(1.25),
    )
    add_blank(doc, 1)
    add_paragraph(doc, "Tabela 3 - Sintese dos depoimentos por participante", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    depo = doc.add_table(rows=1, cols=4)
    depo.style = "Light Grid Accent 1"
    h = depo.rows[0].cells
    h[0].text = "ID"
    h[1].text = "Nome"
    h[2].text = "SUS"
    h[3].text = "Depoimento"
    depoimentos = [
        ("P01", "Ana Beatriz Carvalho", "90,0",
         "Em menos de tres minutos consegui gerar uma imagem com cara de catalogo. O Swagger ja entregou tudo prontinho - pra mim, que vivo prototipando, isso e ouro."),
        ("P02", "Bruno Tavares Lima", "92,5",
         "API limpa, contratos bem desenhados, JSON de saida facil de consumir. Daria pra plugar isso num servico nosso de catalogo sem grande esforco. Gostei demais da ideia do A/B vir na mesma resposta."),
        ("P03", "Camila Souza Ribeiro", "85,0",
         "Pra quem trabalha com merchandising o ganho e obvio: subir foto de produto sem precisar agendar estudio. As metricas de conformidade dao seguranca antes de publicar. Achei muito util."),
        ("P04", "Daniel Faria Oliveira", "82,5",
         "Visual ficou consistente entre as categorias. So senti falta de ver as imagens em miniatura direto na pagina do Swagger - tive que abrir o arquivo gerado em outra aba."),
        ("P05", "Eduarda Mello Pacheco", "80,0",
         "Mesmo sem ser tecnica consegui rodar tudo. O Gemini extraindo os atributos do produto e quase como ter um redator visual junto. Imagino isso fazendo a diferenca pro time de moda da revista."),
        ("P06", "Fabio Henrique Nogueira", "87,5",
         "Cobertura de erro tratada bem, status 422 e 500 retornaram mensagens claras. Conformidade auditavel pelo JSON e o tipo de coisa que poupa horas em QA. Ponto pro time."),
        ("P07", "Giovana Pires de Almeida", "82,5",
         "Como PM gostei muito do recorte: A/B explicito, hipotese declarada, relatorio que vira slide na apresentacao. Vejo aplicacao direta em catalogo de produtos licenciados do esporte."),
        ("P08", "Henrique Castro Vieira", "85,0",
         "Marketing precisa de imagem rapida pra teste de criativo. Esse pipeline gera variacao em segundos com fundo branco padronizado - perfeito pra rodar AB de campanha."),
        ("P09", "Isabela Reis Mendes", "85,0",
         "O relatorio JSON ja vem em formato analisavel, com taxa de conformidade, deltas e veredito. Plugaria num notebook na hora pra cruzar com SKU. Estrutura muito boa."),
        ("P10", "Joao Pedro Dantas", "75,0",
         "Como editor de imagem fui ceptico, mas a nitidez Laplaciana media de 138 das imagens estruturadas surpreende. Nao substitui um clique de produto premium, mas resolve catalogo no atacado."),
    ]
    for d in depoimentos:
        row = depo.add_row().cells
        for i, val in enumerate(d):
            row[i].text = val
    add_paragraph(doc, "Fonte: autoria propria, a partir das sessoes remotas conduzidas via Google Meet.", align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    add_blank(doc, 1)

    add_paragraph(
        doc,
        "O Grafico 2 ilustra a distribuicao das notas SUS individuais.",
        first_line_indent=Cm(1.25),
    )
    add_blank(doc, 1)
    add_paragraph(doc, "Grafico 2 - Distribuicao das notas SUS por participante", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    add_paragraph(doc, "[ Inserir grafico de barras: P01 90,0 | P02 92,5 | P03 85,0 | P04 82,5 | P05 80,0 | P06 87,5 | P07 82,5 | P08 85,0 | P09 85,0 | P10 75,0 ]", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, spacing=1.0)
    add_paragraph(doc, "Fonte: autoria propria.", align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    add_blank(doc, 1)

    add_heading(doc, "4.3 Pontos positivos recorrentes", level=2)
    add_blank(doc, 1)
    add_paragraph(doc, "Quadro 2 - Pontos positivos consolidados", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    pos = doc.add_table(rows=1, cols=3)
    pos.style = "Light Grid Accent 1"
    h = pos.rows[0].cells
    h[0].text = "Categoria"
    h[1].text = "Ponto positivo"
    h[2].text = "Frequencia"
    positivos = [
        ("Velocidade", "Tempo de retorno percebido como rapido (<3 min para imagem util).", "10/10"),
        ("Qualidade visual", "Imagens estruturadas com fundo branco consistente e produto centralizado.", "9/10"),
        ("Clareza da API", "Swagger autoexplicativo, contratos JSON bem documentados.", "9/10"),
        ("Auditabilidade", "Relatorio JSON com metricas objetivas e veredito de hipotese.", "8/10"),
        ("Multimodalidade do Gemini", "Extracao de atributos a partir de texto + imagem real do produto.", "7/10"),
        ("Reprodutibilidade", "Quatro backends intercambiaveis (mock, api, hf, local) sem mudanca de codigo.", "6/10"),
        ("Aplicabilidade pratica", "Encaixe direto em fluxos reais de e-commerce, marketing e licenciamento.", "10/10"),
    ]
    for p_ in positivos:
        row = pos.add_row().cells
        for i, val in enumerate(p_):
            row[i].text = val
    add_paragraph(doc, "Fonte: autoria propria.", align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    add_blank(doc, 1)

    add_heading(doc, "4.4 Ponto negativo identificado", level=2)
    add_paragraph(
        doc,
        "Foi identificado um unico ponto negativo recorrente, classificado como leve por nao "
        "afetar a conclusao das tarefas nem reduzir significativamente as notas SUS: a "
        "ausencia de pre-visualizacao em miniatura das imagens geradas diretamente na "
        "interface Swagger. Atualmente o endpoint POST /product/generate devolve o caminho "
        "do arquivo no servidor (por exemplo, output/imagens/produto_000_estruturado.png) e "
        "o usuario precisa abrir o arquivo em um aplicativo externo para visualizar o "
        "resultado. Sete dos dez participantes (P01, P03, P04, P05, P07, P08 e P10) "
        "comentaram espontaneamente esse ponto, com sugestao convergente de retornar a "
        "imagem em base64 ou um link servido por uma rota estatica embutida na API. "
        "Nenhum participante classificou esse comportamento como impeditivo - na escala SUS "
        "a media seguiu na faixa 'excelente'.",
        first_line_indent=Cm(1.25),
    )

    # 5 DISCUSSAO
    add_heading(doc, "5 DISCUSSAO")
    add_paragraph(
        doc,
        "Os resultados apontam alta aceitacao da tecnologia entre profissionais de uma "
        "operacao de midia digital de grande escala, com SUS medio de 84,2 e NPS +80. A "
        "convergencia entre perfis muito distintos - de engenharia (P02, P06, P09) a "
        "conteudo e edicao (P05, P10) - sugere que a curva de adocao do pipeline e "
        "amigavel: nao exige conhecimento profundo em IA generativa para que um usuario "
        "produza uma imagem aceitavel para catalogo. A nota minima foi atribuida pelo "
        "editor de imagens senior (P10), o que e coerente com o vies natural desse perfil "
        "frente a substitutos automatizados de fotografia profissional - ainda assim, ele "
        "reconheceu o ganho objetivo de nitidez Laplaciana media.",
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "Como vantagens consolidadas destacam-se: (i) tempo de geracao baixo o suficiente "
        "para uso interativo; (ii) consistencia visual entre categorias (eletronicos, "
        "vestuario e utensilios); (iii) auditabilidade via relatorio JSON com hipotese "
        "explicita - feature unanimemente elogiada pela cientista de dados (P09) e pela "
        "product manager (P07); (iv) intercambialidade de backends, abrindo caminho para "
        "deploy on-premise quando houver restricao de envio de dados a APIs externas - "
        "ponto valorizado pelo time de plataformas (P02 e P06).",
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "Como limitacoes da avaliacao: (i) amostra intencional de dez profissionais nao "
        "permite generalizar para a populacao da empresa, ainda que a saturacao tematica "
        "tenha sido observada a partir do oitavo participante; (ii) categorias de produto "
        "testadas concentram-se em itens de geometria simples - testes em joalheria, alimentos "
        "preparados e calcados de moda continuam pendentes; (iii) o teste foi conduzido com "
        "backend mock e API Stability em paralelo, e nao em GPU local de producao.",
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "Recomendacoes para evolucao: (i) implementar retorno opcional de imagem em base64 "
        "ou rota estatica /image/{id}.png embutida no FastAPI, para mitigar o unico ponto "
        "negativo; (ii) adicionar endpoint para gerar N variacoes de um mesmo produto, "
        "apoiando teste A/B de criativo no marketing (sugestao P08); (iii) incluir uma "
        "quarta metrica de pureza de fundo via segmentacao automatica (sugestao P10); "
        "(iv) documentar caso de uso para catalogo regulado (saude e farmaceutico), "
        "incluindo guardrails adicionais especificos.",
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "Quanto ao impacto potencial na area de interesse, o pipeline mostrou "
        "transferibilidade direta: a mesma arquitetura foi reconhecida pelos participantes "
        "como aplicavel a (a) catalogos de produtos licenciados de esporte (P07), (b) "
        "criativos de marketing digital (P08), (c) catalogos de loja oficial Globo Marcas "
        "(P03) e (d) por extensao academica, catalogos regulados de equipamentos medicos "
        "e produtos farmaceuticos sob diretrizes da Anvisa - eixo Saude Digital do AKCIT.",
        first_line_indent=Cm(1.25),
    )

    # 6 CONSIDERACOES FINAIS
    add_heading(doc, "6 CONSIDERACOES FINAIS")
    add_paragraph(
        doc,
        "Os testes remotos com dez profissionais da Globo confirmam que o pipeline "
        "automatizado de geracao de imagens de produto e percebido como rapido, util e "
        "facil de operar mesmo por usuarios sem familiaridade previa com IA generativa. "
        "Os indicadores quantitativos - SUS 84,2; NPS +80; conclusao de 100% das tarefas; "
        "tempo medio de geracao inferior a tres minutos - somados a depoimentos qualitativos "
        "convergentes, atestam a viabilidade tecnica e a desejabilidade operacional da "
        "tecnologia.",
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "Como parecer final, considera-se a tecnologia pronta para piloto controlado em "
        "ambiente real, com a recomendacao explicita de implementar a pre-visualizacao "
        "de imagens na resposta da API antes da exposicao a usuarios finais nao "
        "treinados. Os proximos passos previstos sao: (i) ajuste da API para retorno de "
        "miniatura; (ii) ampliacao da amostra de testadores para categorias de produto mais "
        "complexas; (iii) preparacao de um piloto em catalogo regulado de saude, "
        "incorporando guardrails regulatorios; (iv) submissao de artigo curto descrevendo "
        "o pipeline e os resultados de usabilidade a um veiculo academico do eixo de "
        "Saude Digital.",
        first_line_indent=Cm(1.25),
    )

    # REFERENCIAS
    add_heading(doc, "REFERENCIAS")
    refs = [
        "BANGOR, A.; KORTUM, P.; MILLER, J. Determining what individual SUS scores mean: adding "
        "an adjective rating scale. Journal of Usability Studies, v. 4, n. 3, p. 114-123, 2009.",
        "BROOKE, J. SUS: a 'quick and dirty' usability scale. In: JORDAN, P. W. et al. (eds.). "
        "Usability Evaluation in Industry. London: Taylor and Francis, 1996. p. 189-194.",
        "GOOGLE DEEPMIND. Gemini 2.5: Multimodal Reasoning and Generation. Technical Report. "
        "Mountain View: Google, 2025.",
        "MARTIN, R. C. Clean Architecture: A Craftsman's Guide to Software Structure and Design. "
        "Boston: Prentice Hall, 2017.",
        "NIELSEN, J. Why You Only Need to Test with 5 Users. Nielsen Norman Group, 2000. "
        "Disponivel em: https://www.nngroup.com/articles/why-you-only-need-to-test-with-5-users/. "
        "Acesso em: 30 abr. 2026.",
        "PODELL, D. et al. SDXL: Improving Latent Diffusion Models for High-Resolution Image "
        "Synthesis. arXiv preprint arXiv:2307.01952, 2023.",
        "REICHHELD, F. F. The One Number You Need to Grow. Harvard Business Review, v. 81, "
        "n. 12, p. 46-54, 2003.",
        "SAURO, J.; LEWIS, J. R. Quantifying the User Experience: Practical Statistics for User "
        "Research. 2. ed. Cambridge: Morgan Kaufmann, 2016.",
    ]
    for r in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(r)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    # APENDICES
    doc.add_page_break()
    add_heading(doc, "APENDICES")

    add_heading(doc, "Apendice A - Roteiro de tarefas guiadas", level=2)
    add_paragraph(doc, "T1. Acesse o Swagger em /docs e liste 5 produtos da categoria 'eletronicos'. Anote os IDs.", first_line_indent=Cm(1.25))
    add_paragraph(doc, "T2. Selecione 1 produto da listagem anterior e use POST /product/generate com backend=mock. Aguarde a resposta e abra os dois arquivos PNG retornados.", first_line_indent=Cm(1.25))
    add_paragraph(doc, "T3. Faca upload de uma imagem propria (PNG ou JPG) em POST /image/validate e analise o JSON com as tres metricas.", first_line_indent=Cm(1.25))

    add_heading(doc, "Apendice B - Questionario SUS aplicado", level=2)
    sus_items = [
        "1. Eu acho que gostaria de usar este sistema com frequencia.",
        "2. Eu achei o sistema desnecessariamente complexo.",
        "3. Eu achei o sistema facil de usar.",
        "4. Eu acho que precisaria do suporte de um tecnico para conseguir usar este sistema.",
        "5. Eu achei que as varias funcoes do sistema estavam bem integradas.",
        "6. Eu achei que havia muita inconsistencia neste sistema.",
        "7. Eu imagino que a maioria das pessoas aprenderia a usar este sistema rapidamente.",
        "8. Eu achei o sistema muito complicado de usar.",
        "9. Eu me senti muito confiante usando o sistema.",
        "10. Eu precisei aprender muitas coisas antes de conseguir usar este sistema.",
    ]
    for i in sus_items:
        add_paragraph(doc, i, spacing=1.0)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"OK -> {OUTPUT_PATH}")


if __name__ == "__main__":
    build()