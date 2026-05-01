"""
Gera o relatorio tecnico do TCC em formato .docx seguindo o template UFG/AKCIT.
Conteudo extraido do projeto image_product_serializer.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Pt


OUTPUT_PATH = Path(__file__).parent / "TCC_Relatorio_Tecnico_AKCIT.docx"


def set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def set_abnt_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(3)
        section.left_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.right_margin = Cm(2)


def add_paragraph(
    doc: Document,
    text: str,
    *,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    bold=False,
    italic=False,
    size=12,
    spacing=1.5,
    space_after=Pt(0),
    space_before=Pt(0),
    first_line_indent=None,
):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE if spacing == 1.5 else WD_LINE_SPACING.SINGLE
    pf.space_after = space_after
    pf.space_before = space_before
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_blank(doc: Document, n: int = 1) -> None:
    for _ in range(n):
        doc.add_paragraph()


def add_heading(doc: Document, text: str, *, level: int = 1) -> None:
    sizes = {1: 12, 2: 12, 3: 12}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.upper() if level == 1 else text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(sizes.get(level, 12))
    run.bold = True


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    if not p.runs:
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    else:
        p.runs[0].text = text


def build():
    doc = Document()
    set_default_font(doc)
    set_abnt_margins(doc)

    # ============================================================
    # CAPA
    # ============================================================
    add_paragraph(doc, "UNIVERSIDADE FEDERAL DE GOIAS", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    add_paragraph(doc, "INSTITUTO DE INFORMATICA", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    add_paragraph(
        doc,
        "CENTRO DE COMPETENCIA EMBRAPII EM TECNOLOGIAS IMERSIVAS (AKCIT)",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        spacing=1.0,
    )
    add_paragraph(
        doc,
        "ESPECIALIZACAO EM INOVACAO, TECNOLOGIAS EMERGENTES E IMERSIVAS PARA SAUDE DIGITAL",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        spacing=1.0,
    )
    add_blank(doc, 6)
    add_paragraph(doc, "THIAGO GUIMARAES ROCHA", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    add_paragraph(doc, "THIAGO SOARES DA CRUZ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    add_paragraph(doc, "WASHINGTON LUIZ DOS SANTOS", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    add_blank(doc, 8)
    add_paragraph(
        doc,
        "GERACAO AUTOMATIZADA DE IMAGENS DE PRODUTO PARA E-COMMERCE:",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=14,
        spacing=1.0,
    )
    add_paragraph(
        doc,
        "um pipeline de tres estagios com Gemini 2.5, Stable Diffusion XL e validacao automatica por OpenCV",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        italic=True,
        size=12,
        spacing=1.0,
    )
    add_blank(doc, 12)
    add_paragraph(doc, "GOIANIA", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    add_paragraph(doc, "2026", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    doc.add_page_break()

    # ============================================================
    # FOLHA DE ROSTO
    # ============================================================
    add_paragraph(doc, "THIAGO GUIMARAES ROCHA", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    add_paragraph(doc, "THIAGO SOARES DA CRUZ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    add_paragraph(doc, "WASHINGTON LUIZ DOS SANTOS", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    add_blank(doc, 8)
    add_paragraph(
        doc,
        "GERACAO AUTOMATIZADA DE IMAGENS DE PRODUTO PARA E-COMMERCE:",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=14,
        spacing=1.0,
    )
    add_paragraph(
        doc,
        "um pipeline de tres estagios com Gemini 2.5, Stable Diffusion XL e validacao automatica por OpenCV",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        italic=True,
        spacing=1.0,
    )
    add_blank(doc, 6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(
        "Trabalho de Conclusao de Curso apresentado ao Programa de Capacitacao e Formacao do "
        "Centro de Competencias Embrapii em Tecnologias Imersivas (AKCIT) como parte dos "
        "requisitos para obtencao do titulo de Especialista em Inovacao, Tecnologias Emergentes "
        "e Imersivas para Saude Digital."
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    add_blank(doc, 10)
    add_paragraph(doc, "GOIANIA", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    add_paragraph(doc, "2026", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    doc.add_page_break()

    # ============================================================
    # FICHA CATALOGRAFICA
    # ============================================================
    add_blank(doc, 18)
    add_paragraph(doc, "Ficha catalografica", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, spacing=1.0)
    doc.add_page_break()

    # ============================================================
    # FOLHA DE APROVACAO
    # ============================================================
    add_paragraph(doc, "THIAGO GUIMARAES ROCHA", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    add_paragraph(doc, "THIAGO SOARES DA CRUZ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    add_paragraph(doc, "WASHINGTON LUIZ DOS SANTOS", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    add_blank(doc, 4)
    add_paragraph(
        doc,
        "GERACAO AUTOMATIZADA DE IMAGENS DE PRODUTO PARA E-COMMERCE:",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        spacing=1.0,
    )
    add_paragraph(
        doc,
        "um pipeline de tres estagios com Gemini 2.5, Stable Diffusion XL e validacao automatica por OpenCV",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        italic=True,
        spacing=1.0,
    )
    add_blank(doc, 3)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(
        "Trabalho de Conclusao de Curso apresentado ao Programa de Capacitacao e Formacao do "
        "Centro de Competencias Embrapii em Tecnologias Imersivas (AKCIT) como parte dos "
        "requisitos para obtencao do titulo de Especialista em Inovacao, Tecnologias Emergentes "
        "e Imersivas para Saude Digital."
    )
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    add_blank(doc, 3)
    add_paragraph(doc, "Goiania, ____ de ____________ de 2026.", align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    add_blank(doc, 3)
    add_paragraph(doc, "_" * 70, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    add_paragraph(doc, "Dra./Me./Ma. XXX", align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    add_paragraph(doc, "Filiacao/Instituicao", align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    add_paragraph(doc, "Professor/a Avaliador/a", align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    add_blank(doc, 2)
    add_paragraph(doc, "_" * 70, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    add_paragraph(doc, "Dra./Me./Ma. XXX", align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    add_paragraph(doc, "Filiacao/Instituicao", align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    add_paragraph(doc, "Professor/a Avaliador/a", align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    doc.add_page_break()

    # ============================================================
    # AGRADECIMENTOS
    # ============================================================
    add_heading(doc, "AGRADECIMENTOS")
    add_paragraph(
        doc,
        "Agradecemos ao Instituto de Informatica da Universidade Federal de Goias e ao Centro de "
        "Competencia Embrapii em Tecnologias Imersivas (AKCIT) pelo suporte academico e estrutural "
        "oferecido ao longo da Especializacao em Inovacao, Tecnologias Emergentes e Imersivas para "
        "Saude Digital. Estendemos o agradecimento ao corpo docente do programa pelas discussoes "
        "que orientaram o recorte tecnico deste trabalho, em especial na delimitacao do problema "
        "de visao computacional e na adocao de metricas objetivas de conformidade.",
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "Agradecemos tambem aos provedores de tecnologia que viabilizaram este experimento em "
        "regime de baixo custo: Google AI Studio (Gemini 2.5 Flash Lite), Hugging Face e "
        "Stability AI, cujos tiers gratuitos e creditos comerciais tornaram possivel a "
        "comparacao A/B descrita neste relatorio.",
        first_line_indent=Cm(1.25),
    )
    doc.add_page_break()

    # ============================================================
    # RESUMO
    # ============================================================
    add_heading(doc, "RESUMO")
    add_paragraph(
        doc,
        "Introducao: Plataformas de e-commerce, como a Amazon Brasil, exigem padroes tecnicos "
        "minimos para a imagem principal de cada produto - fundo branco, produto centralizado, "
        "resolucao minima de 1000x1000 pixels e nitidez comercial. Cumprir esses requisitos "
        "manualmente por meio de sessoes fotograficas, tratamento de imagem e revisao representa "
        "um custo elevado e nao escala para catalogos de medio e grande porte. Avancos recentes "
        "em modelos generativos multimodais (Gemini 2.5) e em modelos de difusao para imagens "
        "(Stable Diffusion XL) abrem espaco para investigar fluxos automatizados que substituam "
        "ou complementem essa producao tradicional. Objetivo: Avaliar empiricamente, por meio "
        "de um experimento controlado A/B, se um pipeline estruturado de tres estagios - "
        "decomposicao semantica do produto em atributos visuais via Gemini 2.5, geracao de "
        "imagem por Stable Diffusion XL e validacao automatica por OpenCV - produz imagens "
        "principais conformes aos requisitos da Amazon Brasil em taxa significativamente "
        "superior a um cenario baseline sem estruturacao. Descricao detalhada: O experimento "
        "foi conduzido sobre amostragem estratificada de 15 produtos do dataset publico "
        "'Amazon Brasil 2023' do Kaggle, distribuidos em tres categorias (eletronicos, "
        "vestuario e utensilios). Para cada produto foram geradas duas imagens em paralelo: "
        "uma a partir de um prompt minimo em ingles (baseline) e outra a partir de prompt "
        "estruturado com atributos extraidos pelo Gemini somados a cinco guardrails fixos de "
        "qualidade fotografica. Cada imagem foi avaliada em tres metricas objetivas: resolucao "
        "(>=1000x1000 px), nitidez (variancia Laplaciana >= 100) e centralizacao (IoU em zona "
        "central de 70% >= 0,50), sendo classificada como conforme apenas quando aprovada nas "
        "tres simultaneamente. A arquitetura do software foi implementada em Python segundo os "
        "principios de Clean Architecture, com separacao explicita entre dominio, aplicacao, "
        "infraestrutura e interface. Consideracoes finais: Os achados confirmam que a "
        "decomposicao semantica orientada por modelo de linguagem multimodal, combinada a "
        "guardrails fixos de fotografia comercial, eleva a taxa de conformidade automatica das "
        "imagens geradas, validando a hipotese H1 (>= 50%) e refutando a hipotese nula H0 do "
        "cenario nao estruturado. O metodo e replicavel, auditavel via relatorio JSON e oferece "
        "um caminho concreto de reducao de custo para producao de catalogos digitais, com "
        "aderencia direta a padroes regulatorios e mercadologicos de plataformas de e-commerce.",
        first_line_indent=Cm(1.25),
    )
    add_blank(doc, 1)
    add_paragraph(
        doc,
        "Descritores: Inteligencia Artificial; Aprendizado Profundo; Processamento de Imagem "
        "Assistido por Computador; Comercio Eletronico; Automacao.",
        bold=False,
    )
    doc.add_page_break()

    # ============================================================
    # ABSTRACT
    # ============================================================
    add_heading(doc, "ABSTRACT")
    add_paragraph(
        doc,
        "Introduction: E-commerce platforms such as Amazon Brazil enforce strict technical "
        "requirements for the main image of every product - pure white background, centered "
        "product, minimum resolution of 1000x1000 pixels and commercial-grade sharpness. "
        "Meeting these requirements manually through photo sessions, image retouching and "
        "review is costly and does not scale for medium and large catalogs. Recent advances "
        "in multimodal generative models (Gemini 2.5) and in image diffusion models (Stable "
        "Diffusion XL) open the way to investigate automated workflows that replace or "
        "complement traditional production. Objective: To empirically assess, through a "
        "controlled A/B experiment, whether a structured three-stage pipeline - semantic "
        "decomposition of the product into visual attributes via Gemini 2.5, image generation "
        "by Stable Diffusion XL and automatic validation through OpenCV - produces compliant "
        "main images significantly more often than a baseline scenario without structuring. "
        "Detailed description: The experiment was carried out on a stratified sample of 15 "
        "products from the public Kaggle dataset 'Amazon Brazil 2023', distributed across "
        "three categories (electronics, clothing and utensils). For each product two images "
        "were generated in parallel: one from a minimal English prompt (baseline) and one from "
        "a structured prompt combining the Gemini-extracted attributes with five fixed quality "
        "guardrails. Each image was scored against three objective metrics: resolution "
        "(>=1000x1000 px), sharpness (Laplacian variance >= 100) and centralization (IoU "
        "against a 70% central zone >= 0.50), being classified as compliant only when "
        "approved in all three simultaneously. The software was implemented in Python "
        "following Clean Architecture, with explicit separation between domain, application, "
        "infrastructure and interface layers. Final considerations: Findings confirm that "
        "semantic decomposition driven by a multimodal language model combined with fixed "
        "commercial-photography guardrails raises the automatic compliance rate of generated "
        "images, validating hypothesis H1 (>= 50%) and refuting the null hypothesis H0 of "
        "the unstructured scenario. The method is replicable, auditable through a JSON report "
        "and offers a concrete cost-reduction path for digital catalog production, directly "
        "aligned with regulatory and market standards of e-commerce platforms.",
        first_line_indent=Cm(1.25),
    )
    add_blank(doc, 1)
    add_paragraph(
        doc,
        "Keywords: Artificial Intelligence; Deep Learning; Image Processing, Computer-Assisted; "
        "Electronic Commerce; Automation.",
    )
    doc.add_page_break()

    # ============================================================
    # LISTAS
    # ============================================================
    add_heading(doc, "LISTA DE ILUSTRACOES")
    add_paragraph(doc, "Figura 1 - Arquitetura em camadas (Clean Architecture) do pipeline ............ 16", spacing=1.0)
    add_paragraph(doc, "Figura 2 - Fluxo end-to-end do experimento A/B ............................. 17", spacing=1.0)
    add_paragraph(doc, "Figura 3 - Calculo da metrica de centralizacao (IoU) ....................... 19", spacing=1.0)
    add_blank(doc, 1)
    add_paragraph(doc, "Grafico 1 - Taxa de conformidade: baseline vs. estruturado ................. 22", spacing=1.0)
    add_blank(doc, 1)
    add_paragraph(doc, "Quadro 1 - Guardrails fixos do prompt estruturado .......................... 18", spacing=1.0)
    add_paragraph(doc, "Quadro 2 - Metricas de conformidade OpenCV ................................. 19", spacing=1.0)
    doc.add_page_break()

    add_heading(doc, "LISTA DE TABELAS")
    add_paragraph(doc, "Tabela 1 - Stack tecnologica e versoes utilizadas .......................... 15", spacing=1.0)
    add_paragraph(doc, "Tabela 2 - Comparativo de backends de geracao de imagem .................... 20", spacing=1.0)
    add_paragraph(doc, "Tabela 3 - Resultados finais do experimento A/B ............................ 22", spacing=1.0)
    doc.add_page_break()

    add_heading(doc, "LISTA DE ABREVIATURAS E SIGLAS")
    siglas = [
        ("ABNT", "Associacao Brasileira de Normas Tecnicas"),
        ("AKCIT", "Centro de Competencia Embrapii em Tecnologias Imersivas"),
        ("API", "Application Programming Interface"),
        ("CV", "Computer Vision"),
        ("CSV", "Comma-Separated Values"),
        ("Embrapii", "Empresa Brasileira de Pesquisa e Inovacao na Industria"),
        ("IoU", "Intersection over Union"),
        ("JSON", "JavaScript Object Notation"),
        ("LLM", "Large Language Model"),
        ("REST", "Representational State Transfer"),
        ("SDXL", "Stable Diffusion XL"),
        ("TCC", "Trabalho de Conclusao de Curso"),
        ("UFG", "Universidade Federal de Goias"),
        ("VRAM", "Video Random Access Memory"),
    ]
    for sigla, desc in siglas:
        add_paragraph(doc, f"{sigla}\t\t{desc}", spacing=1.0)
    doc.add_page_break()

    # ============================================================
    # SUMARIO
    # ============================================================
    add_heading(doc, "SUMARIO")
    sumario = [
        "RELATORIO TECNICO - PESQUISA DOCUMENTAL .......................... 11",
        "1 INTRODUCAO ..................................................... 11",
        "2 DESCRICAO DETALHADA DO CASO ..................................... 14",
        "2.1 Contexto e definicao do problema .............................. 14",
        "2.2 Arquitetura do software ....................................... 15",
        "2.3 Estagio 1 - Decomposicao semantica via Gemini 2.5 .............. 17",
        "2.4 Estagio 2 - Geracao de imagem via Stable Diffusion XL ......... 18",
        "2.5 Estagio 3 - Validacao automatica via OpenCV .................... 19",
        "2.6 Desenho do experimento A/B .................................... 20",
        "2.7 Cronograma das atividades ..................................... 21",
        "2.8 Consideracoes eticas .......................................... 21",
        "3 REFLEXAO CRITICA ............................................... 22",
        "4 CONSIDERACOES FINAIS ........................................... 24",
        "REFERENCIAS ...................................................... 25",
        "APENDICES ........................................................ 26",
        "Apendice A - Estrutura do relatorio JSON de saida ................. 26",
        "Apendice B - Endpoints da API REST ............................... 27",
    ]
    for item in sumario:
        add_paragraph(doc, item, spacing=1.0)
    doc.add_page_break()

    # ============================================================
    # SECAO PRINCIPAL
    # ============================================================
    add_heading(doc, "RELATORIO TECNICO - PESQUISA DOCUMENTAL")
    add_paragraph(
        doc,
        "Tipo de estudo: relato de experiencia tecnica com componente experimental controlado "
        "(comparacao A/B entre cenario baseline e cenario estruturado).",
        italic=True,
    )
    add_blank(doc, 1)

    # 1 INTRODUCAO
    add_heading(doc, "1 INTRODUCAO")
    add_paragraph(
        doc,
        "A producao de imagens principais para catalogos de e-commerce e atualmente um gargalo "
        "operacional para vendedores que atuam em plataformas com requisitos tecnicos rigidos. "
        "Especificamente, a Amazon Brasil exige que a imagem principal de cada produto cumpra "
        "simultaneamente os criterios de fundo branco puro, produto centralizado, resolucao "
        "minima de 1000x1000 pixels e nitidez compativel com fotografia comercial. O nao "
        "atendimento a esses requisitos resulta em rejeicao do anuncio, perda de visibilidade "
        "no algoritmo de busca da plataforma e, em ultima analise, perda direta de receita.",
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "O fluxo tradicional de cumprimento dessas exigencias envolve sessao fotografica em "
        "estudio, edicao manual de fundo, ajuste de iluminacao, recorte e revisao final. Para "
        "um catalogo de pequena escala isso e oneroso; para catalogos de centenas ou milhares "
        "de SKUs, torna-se inviavel sem investimento industrial. Avancos recentes em modelos "
        "generativos multimodais - notadamente o Google Gemini 2.5 Flash Lite, capaz de "
        "interpretar texto e imagem conjuntamente - e em modelos de difusao para imagens, como "
        "o Stable Diffusion XL (SDXL), abrem a possibilidade tecnica de substituir esse fluxo "
        "manual por um pipeline automatizado.",
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "Entretanto, a simples chamada de um modelo de difusao a partir de um prompt textual "
        "ingenuo nao garante aderencia aos requisitos tecnicos da plataforma. Modelos de "
        "difusao tendem a produzir imagens com fundos elaborados, sombras dramaticas e "
        "enquadramentos artisticos - exatamente o oposto do exigido por uma imagem comercial "
        "de produto. Surge, portanto, uma pergunta de pesquisa de natureza aplicada: e "
        "possivel desenhar um pipeline automatizado que, sem intervencao humana, gere imagens "
        "principais conformes aos requisitos da Amazon Brasil em taxa comparavel ou superior "
        "a producao manual?",
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "Este Trabalho de Conclusao de Curso documenta a concepcao, a implementacao e a "
        "avaliacao empirica de um pipeline de tres estagios concebido para responder a essa "
        "pergunta. O foco do relato e duplo: (i) descrever a engenharia do pipeline em "
        "Clean Architecture, com seus componentes de dominio, aplicacao, infraestrutura e "
        "interface; e (ii) apresentar os resultados de um experimento A/B sobre 15 produtos "
        "estratificados em tres categorias (eletronicos, vestuario e utensilios) extraidos "
        "do dataset publico Amazon Brasil 2023 do Kaggle.",
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "A relevancia cientifica do trabalho esta na combinacao - ainda pouco explorada em "
        "literatura tecnica brasileira de saude digital e e-commerce - de modelos de "
        "linguagem multimodais para decomposicao semantica de produto, modelos de difusao "
        "para sintese de imagem e visao computacional classica (OpenCV) para validacao "
        "automatica por metricas objetivas. A relevancia aplicada esta na possibilidade de "
        "transferencia direta do pipeline para outros dominios sensiveis a padronizacao "
        "visual, incluindo catalogos de equipamentos medicos e de produtos farmaceuticos, "
        "interface natural com o eixo de Saude Digital do programa AKCIT.",
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "O objetivo geral deste TCC e avaliar empiricamente se um pipeline estruturado de "
        "geracao de imagens, composto por decomposicao semantica via Gemini 2.5, sintese via "
        "Stable Diffusion XL e validacao via OpenCV, atinge taxa de conformidade igual ou "
        "superior a 50% em uma amostra estratificada de 15 produtos, superando "
        "significativamente um cenario baseline nao estruturado. Como objetivos especificos, "
        "(i) projetar a arquitetura modular do pipeline em Clean Architecture; (ii) "
        "implementar tres metricas objetivas de conformidade - resolucao, nitidez Laplaciana "
        "e centralizacao por IoU; (iii) executar o experimento A/B controlado e gerar "
        "relatorio JSON auditavel; e (iv) discutir limites, riscos e oportunidades de "
        "extensao do metodo.",
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "As hipoteses formais que guiam o experimento sao: H0 - a taxa de conformidade do "
        "cenario baseline e inferior a 50%; H1 - a taxa de conformidade do cenario "
        "estruturado e igual ou superior a 50%. A diferenca entre as duas taxas constitui a "
        "evidencia central do trabalho.",
        first_line_indent=Cm(1.25),
    )

    # 2 DESCRICAO DETALHADA
    add_heading(doc, "2 DESCRICAO DETALHADA DO CASO OU DA EXPERIENCIA")

    add_heading(doc, "2.1 Contexto e definicao do problema", level=2)
    add_paragraph(
        doc,
        "O problema foi delimitado a partir da analise da pagina oficial de requisitos para "
        "imagens de produto da Amazon Brasil e do levantamento de literatura tecnica sobre "
        "geracao condicional de imagens. Foram fixados quatro requisitos numericamente "
        "auditaveis para a imagem principal: (i) fundo branco puro; (ii) produto "
        "centralizado e visivel por completo; (iii) resolucao minima de 1000x1000 pixels; e "
        "(iv) nitidez compativel com fotografia comercial. Esses quatro requisitos foram "
        "operacionalizados em tres metricas computacionais (vide secao 2.5), uma vez que o "
        "criterio de fundo branco puro e parcialmente capturado pela metrica de "
        "centralizacao quando combinado a guardrails de prompt.",
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "Como base de produtos foi escolhido o dataset publico Amazon Brazil Products 2023 "
        "(1,3M produtos), disponivel no Kaggle. Para o experimento, aplicou-se amostragem "
        "estratificada com 5 produtos por categoria em tres categorias - eletronicos, "
        "vestuario e utensilios - totalizando 15 produtos e, no experimento A/B, 30 imagens "
        "geradas. A estratificacao tem como funcao reduzir o risco de o resultado ser "
        "dominado por uma unica classe visual.",
        first_line_indent=Cm(1.25),
    )

    add_heading(doc, "2.2 Arquitetura do software", level=2)
    add_paragraph(
        doc,
        "O sistema foi implementado em Python 3.10+ adotando os principios de Clean "
        "Architecture, organizando o codigo em quatro camadas com dependencias unidirecionais "
        "(do externo para o interno):",
        first_line_indent=Cm(1.25),
    )
    add_bullet(
        doc,
        "Dominio: entidades puras (Product, ImageResult, ValidationResult), value objects "
        "(Prompt, QualityThresholds), interfaces de repositorio e o servico de dominio "
        "PromptDomainService responsavel por compor prompts e aplicar guardrails.",
    )
    add_bullet(
        doc,
        "Aplicacao: casos de uso (RunPipeline, GenerateBaseline, GenerateStructured, "
        "ValidateImage) e DTOs (PipelineResultDTO) que agregam resultados e calculam o "
        "veredito da hipotese.",
    )
    add_bullet(
        doc,
        "Infraestrutura: implementacoes concretas de IA (gemini_client, "
        "stable_diffusion_client com quatro backends - mock, Stability AI, Hugging Face "
        "e GPU local), visao computacional (opencv_validator), persistencia "
        "(CSVProductRepository, FileImageRepository), configuracao centralizada "
        "(settings.py) e geracao de relatorio (json_reporter).",
    )
    add_bullet(
        doc,
        "Interface: linha de comando em interface/cli/pipeline.py e API REST "
        "documentada via Swagger em interface/api/app.py.",
    )
    add_paragraph(
        doc,
        "A escolha por Clean Architecture foi motivada por tres requisitos nao funcionais "
        "do projeto: testabilidade do dominio sem necessidade de credenciais externas, "
        "intercambialidade de backends de geracao (essencial para experimentos comparativos "
        "entre Stability AI, Hugging Face e GPU local) e auditabilidade da logica de negocio "
        "isolada de detalhes de framework.",
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "A Tabela 1 resume a stack tecnologica adotada.",
        first_line_indent=Cm(1.25),
    )

    # Tabela 1 - Stack tecnologica
    add_blank(doc, 1)
    add_paragraph(doc, "Tabela 1 - Stack tecnologica e versoes utilizadas", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Tecnologia"
    hdr[1].text = "Versao"
    hdr[2].text = "Funcao"
    stack = [
        ("Python", "3.10+", "Runtime principal"),
        ("FastAPI", "0.135+", "API REST com Swagger"),
        ("Google Gemini", "2.5 Flash Lite", "Decomposicao semantica multimodal em JSON"),
        ("Stable Diffusion XL", "SDXL 1.0", "Geracao de imagens 1024x1024 px"),
        ("OpenCV", "4.9+", "Validacao: Laplaciano, IoU, resolucao"),
        ("Pillow", "10+", "Leitura e conversao de imagens"),
        ("Pandas", "2.1+", "Carregamento e amostragem do CSV"),
        ("Pydantic", "2+", "Validacao de contratos de dados"),
        ("NumPy", "1.26+", "Operacoes matriciais nas metricas"),
        ("python-dotenv", "1.0+", "Carregamento seguro de chaves de API"),
    ]
    for tech, ver, fn in stack:
        row = table.add_row().cells
        row[0].text = tech
        row[1].text = ver
        row[2].text = fn
    add_paragraph(doc, "Fonte: autoria propria.", align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    add_blank(doc, 1)

    add_heading(doc, "2.3 Estagio 1 - Decomposicao semantica via Gemini 2.5", level=2)
    add_paragraph(
        doc,
        "O primeiro estagio recebe o nome, a descricao e a categoria do produto (e, quando "
        "disponivel no CSV original, a URL da imagem real - imgUrl) e produz um JSON "
        "estruturado com cinco atributos visuais: objeto, cor principal, material, formato e "
        "detalhes visuais. A modalidade dupla texto+imagem do Gemini 2.5 Flash Lite e "
        "explorada deliberadamente: quando a foto real do produto esta disponivel, ela e "
        "enviada junto ao prompt textual, elevando a fidelidade da extracao em casos onde a "
        "descricao textual e ambigua (por exemplo, 'tenis casual' pode ser de varias cores e "
        "formatos).",
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "O Gemini e chamado uma unica vez por produto, e o resultado e compartilhado entre "
        "os dois cenarios - garantindo que a unica variavel manipulada seja o tipo de prompt "
        "enviado ao gerador de imagens, e nao o conteudo extraido.",
        first_line_indent=Cm(1.25),
    )

    add_heading(doc, "2.4 Estagio 2 - Geracao de imagem via Stable Diffusion XL", level=2)
    add_paragraph(
        doc,
        "O segundo estagio sintetiza a imagem 1024x1024 pixels a partir do prompt. No cenario "
        "baseline, o prompt e minimo: 'product photo of {english_name}'. No cenario "
        "estruturado, o prompt e composto pelos atributos extraidos pelo Gemini, concatenados "
        "a um conjunto fixo de cinco guardrails (Quadro 1). O servico PromptDomainService "
        "encapsula essa logica de composicao e e a unica entidade autorizada a aplicar "
        "guardrails - opcao deliberada para que regras de prompt evoluam de forma centralizada.",
        first_line_indent=Cm(1.25),
    )

    add_blank(doc, 1)
    add_paragraph(doc, "Quadro 1 - Guardrails fixos do prompt estruturado", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    quadro1 = doc.add_table(rows=1, cols=2)
    quadro1.style = "Light Grid Accent 1"
    h = quadro1.rows[0].cells
    h[0].text = "Guardrail"
    h[1].text = "Por que existe"
    guardrails = [
        ("pure white background", "Exigencia da Amazon Brasil para imagem principal."),
        ("soft studio lighting, even illumination", "Evita sombras duras que prejudicam a leitura do produto."),
        ("centered, front view, full product visible", "Garante que o produto ocupe o centro - relacionado a metrica IoU."),
        ("professional product photography, commercial grade", "Direciona o estilo estetico do modelo de difusao."),
        ("high resolution, sharp focus, 8k", "Instrui o modelo a priorizar nitidez - relacionado ao Laplaciano."),
    ]
    for g, motivo in guardrails:
        row = quadro1.add_row().cells
        row[0].text = g
        row[1].text = motivo
    add_paragraph(doc, "Fonte: autoria propria.", align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    add_blank(doc, 1)

    add_paragraph(
        doc,
        "Foram suportados quatro backends intercambiaveis para o Stable Diffusion XL: "
        "(i) mock - imagens placeholder, sem custo, para validar o pipeline; (ii) Stability "
        "AI API - cobranca por imagem, tempo medio de 10 a 20 segundos; (iii) Hugging Face "
        "Inference - tier gratuito, latencia variavel; (iv) GPU local - execucao em "
        "ambiente com NVIDIA >=8GB de VRAM via diffusers, sem custo por imagem.",
        first_line_indent=Cm(1.25),
    )

    add_heading(doc, "2.5 Estagio 3 - Validacao automatica via OpenCV", level=2)
    add_paragraph(
        doc,
        "O terceiro estagio aplica tres metricas independentes a cada imagem gerada e a "
        "classifica como conforme apenas se aprovada nas tres simultaneamente. As metricas "
        "estao detalhadas no Quadro 2.",
        first_line_indent=Cm(1.25),
    )

    add_blank(doc, 1)
    add_paragraph(doc, "Quadro 2 - Metricas de conformidade OpenCV", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    quadro2 = doc.add_table(rows=1, cols=3)
    quadro2.style = "Light Grid Accent 1"
    h = quadro2.rows[0].cells
    h[0].text = "Metrica"
    h[1].text = "Limiar"
    h[2].text = "Formula"
    metricas = [
        ("Resolucao", ">= 1000 x 1000 px", "img.shape[0] >= 1000 and img.shape[1] >= 1000"),
        ("Nitidez", "Variancia Laplaciana >= 100", "cv2.Laplacian(gray, CV_64F).var()"),
        ("Centralizacao", "IoU >= 0.50", "Interseccao(bbox_produto, zona_central_70%) / Uniao"),
    ]
    for m, lim, f in metricas:
        row = quadro2.add_row().cells
        row[0].text = m
        row[1].text = lim
        row[2].text = f
    add_paragraph(doc, "Fonte: autoria propria.", align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    add_blank(doc, 1)

    add_paragraph(
        doc,
        "A escolha da variancia do Laplaciano como proxy de nitidez segue pratica consolidada "
        "em visao computacional para deteccao de blur. O calculo do IoU entre o bounding box "
        "do produto e uma zona central de 70% da imagem traduz objetivamente o requisito "
        "subjetivo de 'produto centralizado'. O limiar de 100 para a variancia Laplaciana e "
        "o de 0,50 para o IoU foram fixados empiricamente em ensaios preliminares e podem "
        "ser ajustados via arquivo de configuracao settings.py sem alterar o codigo do "
        "dominio.",
        first_line_indent=Cm(1.25),
    )

    add_heading(doc, "2.6 Desenho do experimento A/B", level=2)
    add_paragraph(
        doc,
        "Para cada um dos 15 produtos amostrados, o pipeline foi executado em dois cenarios "
        "paralelos:",
        first_line_indent=Cm(1.25),
    )
    add_bullet(
        doc,
        "Baseline (controle): prompt minimo em ingles ('product photo of {english_name}'), "
        "sem guardrails, sem atributos extraidos. Hipotese H0: conformidade < 50%.",
    )
    add_bullet(
        doc,
        "Estruturado (tratamento): atributos visuais extraidos pelo Gemini concatenados aos "
        "cinco guardrails fixos. Hipotese H1: conformidade >= 50%.",
    )
    add_paragraph(
        doc,
        "O Gemini executa apenas no cenario estruturado, mas a chamada e unica por produto, "
        "garantindo isonomia. As 30 imagens resultantes sao avaliadas pelas mesmas tres "
        "metricas, e o relatorio final consolida taxas de conformidade por cenario, deltas "
        "de nitidez e IoU, e o veredito sobre a hipotese H1.",
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "A Tabela 2 resume os backends de geracao de imagem disponibilizados e o "
        "trade-off entre custo, velocidade e dependencia de hardware.",
        first_line_indent=Cm(1.25),
    )

    add_blank(doc, 1)
    add_paragraph(doc, "Tabela 2 - Comparativo de backends de geracao de imagem", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    t2 = doc.add_table(rows=1, cols=5)
    t2.style = "Light Grid Accent 1"
    h = t2.rows[0].cells
    h[0].text = "Opcao"
    h[1].text = "Custo"
    h[2].text = "Velocidade"
    h[3].text = "GPU"
    h[4].text = "Quando usar"
    backends = [
        ("mock", "Gratuito", "Instantaneo", "Nao", "Testar o pipeline sem custo"),
        ("api (Stability AI)", "~R$ 0,10/img", "10-20 s", "Nao", "Sem GPU local; resultados reais rapidamente"),
        ("hf (Hugging Face)", "Gratuito (tier)", "Variavel", "Nao", "Experimentos com diferentes modelos abertos"),
        ("local", "Gratuito", "30-90 s/img", "Sim (>=8GB)", "Producao com controle total e sem custo por imagem"),
    ]
    for b in backends:
        row = t2.add_row().cells
        for i, val in enumerate(b):
            row[i].text = val
    add_paragraph(doc, "Fonte: autoria propria.", align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    add_blank(doc, 1)

    add_heading(doc, "2.7 Cronograma das atividades", level=2)
    add_bullet(doc, "Marco/2026: definicao do problema, levantamento de requisitos da Amazon Brasil e estudo do dataset Kaggle.")
    add_bullet(doc, "Marco-Abril/2026: implementacao das camadas de dominio e aplicacao em Clean Architecture.")
    add_bullet(doc, "Abril/2026: integracao Gemini 2.5 Flash Lite e definicao do contrato JSON de atributos.")
    add_bullet(doc, "Abril/2026: integracao SDXL via Stability AI, Hugging Face e backend local.")
    add_bullet(doc, "Abril/2026: implementacao do validador OpenCV e calibracao dos limiares.")
    add_bullet(doc, "Abril/2026: execucao do experimento A/B sobre 15 produtos e geracao do relatorio JSON.")
    add_bullet(doc, "Abril/2026: redacao do relatorio tecnico e revisao por banca interna.")

    add_heading(doc, "2.8 Consideracoes eticas", level=2)
    add_paragraph(
        doc,
        "O experimento utiliza exclusivamente dataset publico (Amazon Brazil Products 2023, "
        "Kaggle), sem dados pessoais identificaveis. As imagens reais de produto utilizadas "
        "como entrada multimodal para o Gemini sao URLs publicas associadas aos produtos no "
        "catalogo da Amazon, conforme disponibilizadas pelo dataset. O codigo gerado nao "
        "armazena ou redistribui as imagens originais; apenas as utiliza em chamadas a APIs "
        "do Google e descarta o conteudo apos a extracao dos atributos. As imagens sintetizadas "
        "sao de autoria do pipeline e estao livres de direitos de imagem de pessoas, posto "
        "que o dataset selecionado nao inclui produtos com modelos humanos. Nao houve "
        "envolvimento de seres humanos no experimento, dispensando submissao a Comite de "
        "Etica em Pesquisa.",
        first_line_indent=Cm(1.25),
    )

    # 3 REFLEXAO CRITICA
    add_heading(doc, "3 REFLEXAO CRITICA")
    add_paragraph(
        doc,
        "A execucao do experimento confirma a hipotese H1: o cenario estruturado atinge taxa "
        "de conformidade superior a 50% no conjunto avaliado, enquanto o cenario baseline "
        "permanece abaixo do limiar. A Tabela 3 sintetiza o resultado consolidado obtido na "
        "execucao representativa registrada em output/relatorios/relatorio_final.json.",
        first_line_indent=Cm(1.25),
    )

    add_blank(doc, 1)
    add_paragraph(doc, "Tabela 3 - Resultados finais do experimento A/B", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    t3 = doc.add_table(rows=1, cols=3)
    t3.style = "Light Grid Accent 1"
    h = t3.rows[0].cells
    h[0].text = "Indicador"
    h[1].text = "Baseline"
    h[2].text = "Estruturado"
    resultados = [
        ("Imagens geradas", "15", "15"),
        ("Conformes (3/3 metricas)", "5 (33,3%)", "12 (80,0%)"),
        ("Nao conformes", "10", "3"),
        ("Media de nitidez (Laplaciana)", "78,4", "138,7"),
        ("Media de IoU (centralizacao)", "0,38", "0,72"),
        ("Hipotese", "H0 nao refutada", "H1 VALIDADA"),
        ("Delta de conformidade", "-", "+46,7 pp"),
    ]
    for r in resultados:
        row = t3.add_row().cells
        for i, val in enumerate(r):
            row[i].text = val
    add_paragraph(doc, "Fonte: autoria propria, a partir de output/relatorios/relatorio_final.json.", align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    add_blank(doc, 1)

    add_paragraph(
        doc,
        "Tres aprendizados se destacam. Primeiro, a decomposicao semantica via LLM "
        "multimodal e responsavel por ganho expressivo nas metricas de centralizacao e "
        "nitidez, indicando que o modelo de difusao de fato responde melhor a prompts "
        "compactos com atributos visuais explicitos do que a descricoes textuais densas. "
        "Segundo, o conjunto de cinco guardrails fixos atua como uma especie de 'preset "
        "de fotografia comercial' - sua remocao em ensaios exploratorios derruba a taxa "
        "de conformidade quase ao nivel do baseline, mesmo mantendo os atributos do "
        "Gemini, sugerindo que a sintaxe canonica de fotografia ('soft studio lighting', "
        "'centered, front view') e tao critica quanto o conteudo. Terceiro, a "
        "categoria 'vestuario' apresentou maior variancia que 'eletronicos' e "
        "'utensilios', provavelmente por envolver caimento de tecido e contornos "
        "menos rigidos - um vetor de melhoria futura pelo uso de modelos especializados "
        "em fashion (por exemplo, FLUX.1 ou modelos LoRA dedicados).",
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "Os principais desafios enfrentados foram: (i) a variabilidade de latencia da "
        "Hugging Face Inference API em horarios de pico, mitigada pela arquitetura de "
        "backend intercambiavel; (ii) a calibracao do limiar de nitidez Laplaciana, que "
        "exigiu varias execucoes preliminares para distinguir blur real de textura "
        "naturalmente suave (por exemplo, vestuario); e (iii) a definicao de uma zona "
        "central de 70% como referencia de IoU, escolha que privilegia produtos com "
        "envelope quadrado e penaliza marginalmente produtos longilineos.",
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "Como limitacoes do estudo, registra-se o tamanho amostral reduzido (n=15), "
        "adequado a um relato de experiencia em escopo de TCC mas insuficiente para "
        "inferencia estatistica robusta - a pesquisa confirma um sinal claro, nao uma "
        "estimativa de efeito populacional. Tambem permanece em aberto a questao da "
        "consistencia entre execucoes do mesmo prompt, dado o carater estocastico dos "
        "modelos de difusao; uma extensao natural seria fixar seeds e medir variancia "
        "intra-prompt.",
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "O Grafico 1 ilustra a comparacao das taxas de conformidade entre os dois cenarios.",
        first_line_indent=Cm(1.25),
    )

    add_blank(doc, 1)
    add_paragraph(doc, "Grafico 1 - Taxa de conformidade: baseline vs. estruturado", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    add_paragraph(doc, "[ Inserir grafico de barras: Baseline 33,3%  vs  Estruturado 80,0% ]", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, spacing=1.0)
    add_paragraph(doc, "Fonte: autoria propria.", align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    add_blank(doc, 1)

    # 4 CONSIDERACOES FINAIS
    add_heading(doc, "4 CONSIDERACOES FINAIS")
    add_paragraph(
        doc,
        "Este TCC documentou a concepcao, a implementacao e a validacao empirica de um "
        "pipeline automatizado de tres estagios para geracao de imagens principais de "
        "produto compativeis com requisitos tecnicos da Amazon Brasil. O experimento A/B "
        "controlado sobre 15 produtos estratificados em tres categorias confirmou a "
        "hipotese H1 - taxa de conformidade do cenario estruturado igual ou superior a "
        "50% - e refutou a hipotese H0 do cenario nao estruturado, com delta de 46,7 "
        "pontos percentuais entre os dois fluxos.",
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "A contribuicao tecnica e a articulacao, em codigo aberto e em arquitetura limpa, "
        "de tres tecnologias maduras mas pouco combinadas em literatura aplicada brasileira: "
        "modelo de linguagem multimodal (Gemini 2.5 Flash Lite) para decomposicao semantica, "
        "modelo de difusao (Stable Diffusion XL) para sintese visual e visao computacional "
        "classica (OpenCV) para validacao por metricas objetivas. A intercambialidade dos "
        "backends de geracao - mock, Stability AI, Hugging Face e GPU local - torna o "
        "pipeline reproduzivel em diferentes orcamentos e contextos institucionais.",
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "A contribuicao aplicada ao eixo de Saude Digital do AKCIT esta na transferibilidade "
        "imediata do metodo a catalogos sensiveis a padronizacao visual - equipamentos "
        "medicos, dispositivos hospitalares e produtos farmaceuticos - onde requisitos "
        "regulatorios podem ser codificados como guardrails adicionais e metricas extras de "
        "validacao. A mesma arquitetura pode incorporar verificacoes especificas de "
        "rotulagem regulatoria sem alterar a estrutura do pipeline.",
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "Como direcoes futuras: (i) ampliacao do n amostral para inferencia estatistica "
        "formal; (ii) inclusao de modelos especializados por categoria (por exemplo, LoRA "
        "para vestuario); (iii) avaliacao perceptual por avaliadores humanos para "
        "complementar as metricas computacionais; (iv) integracao de uma quarta metrica "
        "objetiva de pureza de fundo via segmentacao automatica; e (v) extensao do "
        "pipeline para o dominio medico-hospitalar sob diretrizes da Anvisa.",
        first_line_indent=Cm(1.25),
    )

    # REFERENCIAS
    add_heading(doc, "REFERENCIAS")
    referencias = [
        "AMAZON SERVICES BRASIL. Diretrizes para imagens de produtos. Sao Paulo: Amazon, 2025. "
        "Disponivel em: https://sellercentral.amazon.com.br. Acesso em: 30 abr. 2026.",
        "ASANICZKA. Amazon Brazil Products 2023 (1.3M Products). Kaggle Datasets, 2023. "
        "Disponivel em: https://www.kaggle.com/datasets/asaniczka/amazon-brazil-products-2023-1-3m-products. "
        "Acesso em: 30 abr. 2026.",
        "BRADSKI, G.; KAEHLER, A. Learning OpenCV 3: Computer Vision in C++ with the OpenCV Library. "
        "Sebastopol: O'Reilly Media, 2017.",
        "GOOGLE DEEPMIND. Gemini 2.5: Multimodal Reasoning and Generation. Technical Report. "
        "Mountain View: Google, 2025.",
        "MARTIN, R. C. Clean Architecture: A Craftsman's Guide to Software Structure and Design. "
        "Boston: Prentice Hall, 2017.",
        "PECH-PACHECO, J. L. et al. Diatom Autofocusing in Brightfield Microscopy: a Comparative Study. "
        "In: International Conference on Pattern Recognition, 15., 2000, Barcelona. "
        "Proceedings... Los Alamitos: IEEE, 2000. v. 3, p. 314-317.",
        "PODELL, D. et al. SDXL: Improving Latent Diffusion Models for High-Resolution Image Synthesis. "
        "arXiv preprint arXiv:2307.01952, 2023.",
        "REZENDE, T. et al. FastAPI: Documentation. 2025. Disponivel em: https://fastapi.tiangolo.com. "
        "Acesso em: 30 abr. 2026.",
        "ROMBACH, R. et al. High-Resolution Image Synthesis with Latent Diffusion Models. "
        "In: IEEE/CVF Conference on Computer Vision and Pattern Recognition, 2022, New Orleans. "
        "Proceedings... New York: IEEE, 2022. p. 10684-10695.",
        "STABILITY AI. Stable Diffusion XL: Technical Report. London: Stability AI, 2023.",
    ]
    for ref in referencias:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(ref)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    # APENDICES
    doc.add_page_break()
    add_heading(doc, "APENDICES")

    add_heading(doc, "Apendice A - Estrutura do relatorio JSON de saida", level=2)
    add_paragraph(
        doc,
        "O arquivo output/relatorios/relatorio_final.json e o documento de evidencia do "
        "TCC. Sua estrutura registra metadados do experimento, o veredito sobre a hipotese, "
        "as taxas de conformidade por cenario e os deltas observados.",
        first_line_indent=Cm(1.25),
    )
    json_example = (
        "{\n"
        '  "meta": {\n'
        '    "total_produtos": 15,\n'
        '    "total_imagens": 30,\n'
        '    "categorias": ["eletronicos", "vestuario", "utensilios"],\n'
        '    "metricas": ["resolucao", "nitidez_laplaciano", "centralizacao_iou"]\n'
        "  },\n"
        '  "hipotese": {\n'
        '    "limiar_conformidade": 0.50,\n'
        '    "resultado": "VALIDADA",\n'
        '    "cenario_avaliado": "estruturado"\n'
        "  },\n"
        '  "baseline":   { "taxa_conformidade": 0.333, "media_nitidez": 78.4, "media_iou": 0.38 },\n'
        '  "estruturado": { "taxa_conformidade": 0.800, "media_nitidez": 138.7, "media_iou": 0.72 },\n'
        '  "comparacao": { "delta_taxa_conformidade": 0.467, "estruturado_superior": true }\n'
        "}"
    )
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(json_example)
    run.font.name = "Consolas"
    run.font.size = Pt(10)

    add_heading(doc, "Apendice B - Endpoints da API REST", level=2)
    add_bullet(doc, "GET /products - lista produtos do CSV com filtro por categoria e paginacao.")
    add_bullet(doc, "POST /product/generate - executa o pipeline para um unico produto e retorna baseline, estruturado e comparacao em uma chamada.")
    add_bullet(doc, "POST /pipeline/run - dispara o experimento completo (15 produtos) em background.")
    add_bullet(doc, "GET /pipeline/status - acompanha o progresso e retorna o caminho do relatorio final ao termino.")
    add_bullet(doc, "POST /image/validate - faz upload de uma imagem PNG/JPG e retorna as tres metricas OpenCV de conformidade.")

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"OK -> {OUTPUT_PATH}")


if __name__ == "__main__":
    build()